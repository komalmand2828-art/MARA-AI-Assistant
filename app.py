import streamlit as st
from rag import (
    read_multiple_pdfs,
    create_vector_db,
    count_chunks,
    get_pdf_stats
)
from agents import (
    manager_agent,
    notes_agent,
    quiz_agent
)
from reportlab.pdfgen import canvas
from io import BytesIO
from docx import Document
import speech_recognition as sr
import tempfile
from resume_analyzer import analyze_resume
from auth import *
from database import(
    conn,
    cursor,
    get_memory
)
from export_utils import export_pdf
from workflow import graph
import time
# =====================================
# PAGE CONFIG
# =====================================
st.set_page_config(
    page_title="MARA AI Assistant",
    page_icon="🤖",
    layout="wide"
)
theme = st.sidebar.selectbox(
    "🎨 Theme",
    ["Dark","Light"]
)
# =====================================
# LOGIN SYSTEM
# =====================================

username = st.sidebar.text_input(
    "👤 Username"
)

password = st.sidebar.text_input(
    "🔒 Password",
    type="password"
)

if (
    username != "student"
    or password != "student123"
):

    st.warning(
        "Please Login"
    )

    st.stop()
# =====================================
# LOAD CSS
# =====================================
css_file = (
    "dark.css"
    if theme == "Dark"
    else "light.css"
)
with open(css_file) as css:
    st.markdown(
        f"<style>{css.read()}</style>",
        unsafe_allow_html=True
    )
# =====================================
# SESSION STATE
# =====================================
if "messages" not in st.session_state:
    st.session_state.messages = []
if "memory" not in st.session_state:
    st.session_state.memory=[]
if "vector_db" not in st.session_state:
    st.session_state.vector_db = None
if "current_agent" not in st.session_state:
    st.session_state.current_agent = "None"
if "pdf_text" not in st.session_state:
    st.session_state.pdf_text = ""
if "chunk_count" not in st.session_state:
    st.session_state.chunk_count = 0
if "notes" not in st.session_state:
    st.session_state.notes=""
if "quiz" not in st.session_state:
    st.session_state.quiz=""
if "timeline" not in st.session_state:
    st.session_state.timeline= [] 
if "mcqs" not in st.session_state:
    st.session_state.mcqs=""
if "summary" not in st.session_state:
    st.session_state.summary=""
if "viva" not in st.session_state:
    st.session_state.viva=""  
if "voice_prompt" not in st.session_state:
    st.session_state.voice_prompt=""
if "voice_processed" not in st.session_state:
    st.session_state.voice_processed=False
if "chat_sessions" not in st.session_state:

    st.session_state.chat_sessions = {
        "New Chat": []
    }

if "current_chat" not in st.session_state:

    st.session_state.current_chat = (
        "New Chat"
    )
if "history" not in st.session_state:
    st.session_state.history=[]
if "feedback" not in st.session_state:
    st.session_state.feedback=[]
if "confidence" not in st.session_state:
    st.session_state.confidence=0
if "processing_time" not in st.session_state:
    st.session_state.processing_time = 0
if "agent_usage" not in st.session_state:
    st.session_state.agent_usage = {}
if "language" not in st.session_state:
    st.session_state.language = "English"
if "provider" not in st.session_state:
    st.session_state.provider = "Auto"
if "resume_result" not in st.session_state:
    st.session_state.resume_result = ""
# =====================================
# HEADER
# =====================================
st.markdown(
    """
    <div class='main-title rainbow'>
        🤖 MARA<br>
        <span style="font-size:28px;">
            Enterprise Multi-Agent AI Platform
        </span>
    </div>
    """,
    unsafe_allow_html=True
)
st.markdown(
    """
    AI-Powered Multi-Agent Assistant
    with Gemini + RAG + PDF Intelligence
    """
)
st.divider()
col1, col2, col3, col4 = st.columns(4)

with col1:

    st.metric(
        "Questions",
        len(
            [
                m
                for m in st.session_state.chat_sessions[
                    st.session_state.current_chat
                ]
                if m["role"] == "user"
            ]
        )
    )

with col2:

    st.metric(
        "📚 PDFs",
        0
    )

with col3:

    st.metric(
        "🧠 Chunks",
        st.session_state.chunk_count
    )
with col4:
    st.metric(
        "🤖 Active Agent",
        st.session_state.current_agent
    )

st.info(
    f"""
### 🟢 System Status

🤖 **Active Agent:** {st.session_state.current_agent}

🌍 **Language:** {st.session_state.get("language", "English")}

⚡ **Provider:** {st.session_state.get("provider", "Auto")}

📚 **Knowledge Base:** {"Connected" if st.session_state.vector_db else "Not Loaded"}

💬 **Current Session:** {st.session_state.current_chat}
"""
)
# =====================================
# SIDEBAR
# =====================================

st.sidebar.title("🚀 MARA Dashboard")

# =====================================
# SETTINGS
# =====================================

with st.sidebar.expander("⚙️ Settings", expanded=True):


    language = st.selectbox(
        "🌍 Language",
        [
            "English",
            "Hindi",
            "Punjabi",
            "French",
            "Spanish",
            "German"
        ]
    )

    voice_language = st.selectbox(
        "🎤 Voice Language",
        [
            "English",
            "Hindi",
            "Punjabi"
        ]
    )

    provider = st.selectbox(
        "⚡ AI Provider",
        [
            "Auto",
            "OpenRouter",
            "Gemini"
        ]
    )

    mode = st.selectbox(
        "🎯 Response Mode",
        [
            "Professional",
            "Research",
            "Teacher",
            "Exam Helper"
        ]
    )

st.session_state.language = language
st.session_state.provider = provider

# =====================================
# CHAT
# =====================================

with st.sidebar.expander("💬 Chat", expanded=True):

    if st.button("➕ New Chat",
                 key="new_chat_btn"):

        new_chat = f"Chat {len(st.session_state.chat_sessions)+1}"

        st.session_state.chat_sessions[new_chat] = []

        st.session_state.current_chat = new_chat

        st.rerun()

    selected_chat = st.selectbox(
        "Select Session",
        list(st.session_state.chat_sessions.keys()),
        key="chat_select"
    )

    st.session_state.current_chat = selected_chat

    if st.button("📄 Export Chat PDF",
                 key="export_pdf_btn"):

        pdf_file = export_pdf(st.session_state.messages)

        with open(pdf_file, "rb") as file:

            st.download_button(
                "Download PDF",
                file,
                file_name="MARA_Chat.pdf"
            )

    if st.button("🗑 Clear Chat",
                 key="clear_chat_btn"):

        st.session_state.chat_sessions[
            st.session_state.current_chat
        ] = []

        st.rerun()

# =====================================
# KNOWLEDGE BASE
# =====================================

with st.sidebar.expander("📚 Knowledge Base", expanded=True):

    uploaded_files = st.file_uploader(
        "Upload PDF Documents",
        type=["pdf"],
        accept_multiple_files=True,
        key="pdf_upload"
    )

    web_search_enabled = st.checkbox(
        "🌐 Enable Web Search",
        value=False,
        key="web_search"
    )

    st.subheader("📝 Study Tools")

    # Keep your existing
    # Generate Notes
    # Generate Quiz
    # Generate MCQs
    # Generate Summary
    # buttons here

# =====================================
# RESUME ANALYZER
# =====================================

with st.sidebar.expander("📄 Resume Analyzer"):

    resume_file = st.file_uploader(
        "Upload Resume",
        type=["pdf"],
        key="resume_upload"
    )

    if st.button("Analyze Resume",
                 key="resume_btn"):

        if resume_file:

            with st.spinner("Analyzing Resume..."):

                result = analyze_resume(resume_file)

            st.session_state.resume_result = result

# =====================================
# ANALYTICS
# =====================================

with st.sidebar.expander("📊 Analytics"):

    st.metric(
        "Confidence",
        f"{st.session_state.confidence}%"
    )

    st.metric(
        "Knowledge Chunks",
        st.session_state.chunk_count
    )

    st.metric(
        "Chat Sessions",
        len(st.session_state.chat_sessions)
    )

    st.info(
        f"Current Agent:\n\n{st.session_state.current_agent}"
    )

    st.success(
        f"Stored Memories: {len(get_memory())}"
    )

   

# =====================================
# ABOUT
# =====================================

with st.sidebar.expander("ℹ️ About MARA"):
    st.write("MARA AI Assistant")


# =====================================
# PDF STATS
# =====================================
pdf_count = 0
page_count = 0
if uploaded_files:
    stats = get_pdf_stats(
        uploaded_files
    )
    pdf_count = stats["pdfs"]
    page_count = stats["pages"]
# =====================================
# METRICS
# =====================================

st.sidebar.subheader(
    "📊 Session Metrics"
)
st.sidebar.metric(
    "Agent Confidence",
    f"{st.session_state.confidence}%"
)

st.sidebar.metric(
    "Processing Time",
    f"{st.session_state.processing_time} sec"
)

# Questions in current chat only
st.sidebar.metric(
    "Questions Asked",
    len(
        [
            msg
            for msg in st.session_state.chat_sessions[
                st.session_state.current_chat
            ]
            if msg["role"] == "user"
        ]
    )
)

# PDFs Loaded
st.sidebar.metric(
    "PDFs Loaded",
    pdf_count
)

# Total PDF Pages
st.sidebar.metric(
    "PDF Pages",
    page_count
)

# Knowledge Chunks
st.sidebar.metric(
    "Knowledge Chunks",
    st.session_state.chunk_count
)

# Total Chat Sessions
st.sidebar.metric(
    "Chat Sessions",
    len(
        st.session_state.chat_sessions
    )
)

# Current Session Name
st.sidebar.info(
    f"📂 Active Session:\n\n{st.session_state.current_chat}"
)

# =====================================
# USER FEEDBACK SYSTEM
# =====================================

st.sidebar.markdown("---")
st.sidebar.subheader("⭐ User Feedback")

feedback = st.sidebar.radio(
    "Rate MARA",
    ["⭐ Excellent", "👍 Good", "😐 Average", "👎 Poor"],
    key="feedback_rating"
)

feedback_text = st.sidebar.text_area(
    "Comments",
    key="feedback_text"
)

if st.sidebar.button(
    "Submit Feedback",
    key="submit_feedback_btn"
):
    st.sidebar.success(
        "✅ Thank you for your feedback!"
    )

# =====================================
# KNOWLEDGE STATUS
# =====================================

st.sidebar.subheader(
    "🧠 Knowledge Base"
)

if st.session_state.vector_db:
    st.sidebar.success(
        "🟢 Active"
    )
else:
    st.sidebar.warning(
        "🟡 No PDF Loaded"
    )

# =====================================
# AGENT VISUALIZATION
# =====================================

st.subheader("🧠 MARA LangGraph Workflow")

agent = st.session_state.current_agent

workflow = {
    "General Agent": [
        "👤 User Query",
        "🧭 Router",
        "🧠 Memory",
        "🤖 General Agent",
        "✨ Response"
    ],

    "PDF Agent": [
        "👤 User Query",
        "🧭 Router",
        "🧠 Memory",
        "📚 PDF Agent",
        "🔍 Vector Search",
        "✨ Response"
    ],

    "Study Agent": [
        "👤 User Query",
        "🧭 Router",
        "🧠 Memory",
        "🎓 Study Agent",
        "📖 Learning",
        "✨ Response"
    ],

    "Research Agent": [
        "👤 User Query",
        "🧭 Router",
        "🧠 Memory",
        "🔬 Research Agent",
        "🌐 Web Search",
        "📊 Analysis",
        "✨ Response"
    ],

    "Career Agent": [
        "👤 User Query",
        "🧭 Router",
        "🧠 Memory",
        "💼 Career Agent",
        "🚀 Career Guidance",
        "✨ Response"
    ],

    "Web Agent": [
        "👤 User Query",
        "🧭 Router",
        "🧠 Memory",
        "🌐 Web Agent",
        "🔍 Internet Search",
        "✨ Response"
    ],

    "Document Agent": [
        "👤 User Query",
        "🧭 Router",
        "🧠 Memory",
        "📄 Document Agent",
        "📑 Document Processing",
        "✨ Response"
    ],

    "Email Agent": [
        "👤 User Query",
        "🧭 Router",
        "🧠 Memory",
        "📧 Email Agent",
        "✍️ Email Drafting",
        "✨ Response"
    ],

    "Calendar Agent": [
        "👤 User Query",
        "🧭 Router",
        "🧠 Memory",
        "📅 Calendar Agent",
        "🗓 Task Scheduling",
        "✨ Response"
    ],

    "Coding Agent": [
        "👤 User Query",
        "🧭 Router",
        "🧠 Memory",
        "💻 Coding Agent",
        "⚙️ Code Analysis",
        "✨ Response"
    ],

    "Calculator Tool": [
        "👤 User Query",
        "🧭 Router",
        "🧮 Calculator Tool",
        "✨ Result"
    ],

    "Time Tool": [
        "👤 User Query",
        "🧭 Router",
        "🕒 Time Tool",
        "✨ Result"
    ],

    "Web Tool": [
        "👤 User Query",
        "🧭 Router",
        "🌐 Web Tool",
        "✨ Result"
    ]
}

steps = workflow.get(
    agent,
    workflow["General Agent"]
)

cols = st.columns(len(steps))

for col, step in zip(cols, steps):

    with col:
        st.info(step)

# =====================================
# NOTES GENERATOR
# =====================================

st.sidebar.subheader(
    "📝 Study Tools",

)

# -----------------------------
# GENERATE NOTES
# -----------------------------
if st.sidebar.button(
    "Generate Notes",
    key="notes_btn"
):

    if st.session_state.pdf_text:

        with st.spinner(
            "Generating Notes..."
        ):

            st.session_state.notes = (
                notes_agent(
                    st.session_state.pdf_text[:12000]
                )
            )

    else:

        st.sidebar.warning(
            "Upload PDF First"
        )

# -----------------------------
# GENERATE QUIZ
# -----------------------------
if st.sidebar.button(
    "Generate Quiz",
    key="quiz_btn"
):

    if st.session_state.pdf_text:

        with st.spinner(
            "Generating Quiz..."
        ):

            st.session_state.quiz = (
                quiz_agent(
                    st.session_state.pdf_text[:12000]
                )
            )

    else:

        st.sidebar.warning(
            "Upload PDF First"
        )

# -----------------------------
# GENERATE MCQs
# -----------------------------
if st.sidebar.button(
    "Generate MCQs",
    key="mcq_btn"
):

    if st.session_state.pdf_text:

        with st.spinner(
            "Generating MCQs..."
        ):

            prompt = f"""
            Based on the following PDF content:

            {st.session_state.pdf_text[:12000]}

            Generate:
            - 10 Multiple Choice Questions
            - 4 options for each question
            - Mark the correct answer
            - Make questions exam-oriented
            """

            result = manager_agent(
                query=prompt,
                vector_db=st.session_state.vector_db,
                mode="Teacher",
                history=""
            )

            st.session_state.mcqs = (
                result["response"]
            )

    else:

        st.sidebar.warning(
            "Upload PDF First"
        )

# -----------------------------
# GENERATE SUMMARY
# -----------------------------
if st.sidebar.button(
    "Generate Summary",
    key="summary_btn"
):

    if st.session_state.pdf_text:

        with st.spinner(
            "Generating Summary..."
        ):

            prompt = f"""
            Summarize the following PDF.

            Include:
            - Key Concepts
            - Important Topics
            - Important Definitions
            - Exam Tips
            - Quick Revision Notes

            PDF Content:

            {st.session_state.pdf_text[:12000]}
            """

            result = manager_agent(
                query=prompt,
                vector_db=st.session_state.vector_db,
                mode="Teacher",
                history=""
            )

            st.session_state.summary = (
                result["response"]
            )

    else:

        st.sidebar.warning(
            "Upload PDF First"
        )
# =====================================
# DOWNLOAD NOTES PDF
# =====================================

if st.session_state.notes:

    pdf_buffer = BytesIO()

    c = canvas.Canvas(pdf_buffer)

    text = c.beginText(40, 800)

    for line in st.session_state.notes.split("\n"):

        text.textLine(line)

    c.drawText(text)

    c.save()

    pdf_buffer.seek(0)

    st.sidebar.download_button(
        label="📥 Download Notes PDF",
        data=pdf_buffer,
        file_name="MARA_Notes.pdf",
        mime="application/pdf"
    )
# =====================================
# MEMORY
# =====================================

st.sidebar.subheader(
    "🧠 Long Term Memory"
)

memory_items = get_memory()

st.sidebar.success(
    f"Stored Memories: {len(memory_items)}"
)
# =====================================
# ABOUT
# =====================================
st.sidebar.subheader(
    "ℹ About MARA"
)
st.sidebar.markdown(
    """
    Features:
    ✅ Gemini AI
    ✅ Multi-Agent System
    ✅ PDF Chat
    ✅ RAG
    ✅ FAISS Vector Search
    ✅ Study Assistant
    ✅ Research Assistant
    ✅ Career Mentor
    ✅ Document Agent
    ✅ Email Agent
    ✅ Calendar Agent
    ✅ Coding Agent
    """
)
# =====================================
# DOWNLOAD CHAT
# =====================================

chat_text = ""

for msg in st.session_state.messages:

    role = msg["role"]

    content = msg["content"]

    chat_text += (
        f"{role.upper()}:\n"
        f"{content}\n\n"
    )

st.sidebar.download_button(
    label="📥 Download Chat",
    data=chat_text,
    file_name="MARA_chat.txt",
    mime="text/plain"
)
# =====================================
# DOWNLOAD CHAT DOCX
# =====================================

doc = Document()

doc.add_heading(
    "MARA Conversation",
    level=1
)

for msg in st.session_state.messages:

    doc.add_paragraph(
        f"{msg['role'].upper()}:"
    )

    doc.add_paragraph(
        msg["content"]
    )

doc_buffer = BytesIO()

doc.save(doc_buffer)

doc_buffer.seek(0)

st.sidebar.download_button(
    label="📄 Download DOCX",
    data=doc_buffer,
    file_name="MARA_Chat.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
# =====================================
# STUDY TABS
# =====================================

tab1, tab2, tab3 = st.tabs(
    [
        "💬 Chat",
        "❓ Questions",
        "📖 Summary"
    ]
)

with tab2:

    st.subheader(
        "❓ Generated Questions"
    )

    if st.session_state.mcqs:

        st.markdown(
            st.session_state.mcqs
        )

    if st.session_state.quiz:

        st.markdown(
            st.session_state.quiz
        )

with tab3:

    st.subheader(
        "📖 PDF Summary"
    )

    if st.session_state.summary:

        st.markdown(
            st.session_state.summary
        )
# =====================================
# CHAT HISTORY
# =====================================

with tab1:

    for message in st.session_state.chat_sessions[
    st.session_state.current_chat
]:

        with st.chat_message(
            message["role"]
        ):

            st.markdown(
                message["content"]
            )
# =====================================
# SUMMARY TAB
# =====================================
if "resume_result" in st.session_state:

    with st.expander(
        "📄 Resume Analysis"
    ):

        st.markdown(
            st.session_state.resume_result
        )

# =====================================
# VOICE INPUT
# =====================================

st.caption("🎤 Voice Input")

audio_file = st.audio_input(
    "🎤 Speak to MARA"
)

if audio_file is not None:

    st.caption("🎤 Audio Received")

    try:

        with tempfile.NamedTemporaryFile(
            delete=False,
            suffix=".wav"
        ) as f:

            f.write(audio_file.read())

            audio_path = f.name

        recognizer = sr.Recognizer()

        with sr.AudioFile(audio_path) as source:

            audio = recognizer.record(source)

        recognized_text = (
            recognizer.recognize_google(audio)
        )

        st.success(
            f"Recognized: {recognized_text}"
        )

        st.session_state.voice_prompt = (
            recognized_text
        )

        if (
            not st.session_state.voice_processed
        ):

            st.session_state.voice_processed = True

            st.rerun()

    except Exception as e:

        st.error(
            f"Voice Error: {e}"
        )
# =====================================
# CHAT INPUT
# =====================================
text_prompt = st.chat_input(
    "Ask MARA anything..."
)

prompt = ""

if text_prompt:
    prompt=text_prompt
elif st.session_state.voice_prompt:
    prompt= st.session_state.voice_prompt

# =====================================
# RESPONSE GENERATION
# =====================================

if prompt:

    # -------------------------
    # USER MESSAGE
    # -------------------------

    st.session_state.messages.append(
        {
            "role": "user",
            "content": prompt
        }
    )

    st.session_state.chat_sessions[
        st.session_state.current_chat
    ].append(
        {
            "role": "user",
            "content": prompt
        }
    )

    st.session_state.voice_prompt = ""
    st.session_state.voice_processed = False

    with st.chat_message("user"):
        st.markdown(prompt)

    # -------------------------
    # CHAT HISTORY
    # -------------------------

    history = ""

    for msg in st.session_state.chat_sessions[
        st.session_state.current_chat
    ][-6:]:

        history += (
            f"{msg['role']}: "
            f"{msg['content']}\n"
        )

    # -------------------------
    # CALL LANGGRAPH
    # -------------------------

    try:
        with st.spinner("🧠 MARA is reasoning using LangGraph..."):
            start = time.time()
            result = graph.invoke(
                {
                "query": prompt,
                "vector_db": st.session_state.vector_db,
                "mode": mode,
                "history": history,
                "provider": provider,
                "memory": st.session_state.memory,
                "language": language,
                "web_search": web_search_enabled
                }
            )

            response = result.get(
            "response",
            "⚠️ No response returned."
            )

            agent = result.get(
            "agent",
            "General Agent"
            )

            confidence = result.get(
            "confidence",
            80
            )
            end = time.time()
            processing_time = round(
                end - start,
                2
            )
            st.toast(
                "✅ LangGraph workflow completed successfully."
            )


    except Exception as e:

        response = (
            f"⚠️ Error Occurred\n\n{str(e)}"
        )

        agent = "System"

        confidence = 0
    # ==========================
    # TIMELINE
    # ==========================

    st.session_state.timeline = []

    st.session_state.timeline.append(
        "🧭 Router Node"
    )

    st.session_state.timeline.append(
        "🧠 Memory Node"
    )

    st.session_state.timeline.append(
        f"🤖 {agent}"
    )

    if agent == "PDF Agent":

        st.session_state.timeline.append(
            "📚 Vector Database Retrieval"
        )

    elif agent == "Web Agent":

        st.session_state.timeline.append(
            "🌐 Internet Search"
        )

    elif agent == "Research Agent":

        st.session_state.timeline.append(
            "📊 Research Analysis"
        )

    elif agent == "Coding Agent":

        st.session_state.timeline.append(
            "💻 Code Analysis"
        )

    elif agent == "Career Agent":

        st.session_state.timeline.append(
            "💼 Career Guidance"
        )

    elif agent == "Document Agent":

        st.session_state.timeline.append(
            "📄 Document Processing"
        )

    elif agent == "Email Agent":

        st.session_state.timeline.append(
            "📧 Email Drafting"
        )

    elif agent == "Calendar Agent":

        st.session_state.timeline.append(
            "📅 Schedule Planning"
        )

    st.session_state.timeline.append(
        f"🎯 Confidence: {confidence}%"
    )

    st.session_state.timeline.append(
        "✅ Response Generated"
    )

    # ==========================
    # CURRENT AGENT
    # ==========================

    if "route" in result:

        st.session_state.current_agent = (
        result["route"].title()
        + " Agent"
        )

    else:

        st.session_state.current_agent = agent
    if agent not in st.session_state.agent_usage:

        st.session_state.agent_usage[agent] = 0

    st.session_state.agent_usage[agent] += 1

    st.session_state.confidence = confidence
    st.session_state.processing_time = processing_time
    # ==========================
    # ASSISTANT MESSAGE
    # ==========================

    with st.chat_message("assistant"):

        st.success(
            f"🤖 **{agent}**"
        )
        st.caption(
            f"🎯 Confidence: {confidence}% | ⏱️ {processing_time} sec"
        )
        st.success("✅ Response Generated")
        st.markdown(response)

    # ==========================
    # SAVE RESPONSE
    # ==========================

    st.session_state.messages.append(
        {
            "role": "assistant",
            "content": response
        }
    )

    st.session_state.chat_sessions[
        st.session_state.current_chat
    ].append(
        {
            "role": "assistant",
            "content": response
        }
    )

    st.session_state.voice_prompt = ""
    # ==========================
    # SAVE TO SQL MEMORY
    # ==========================

    cursor.execute(
         """
         INSERT INTO memory
         (question, answer)
         VALUES (?, ?)
        """,
         (
             prompt,
             response
        )
    )

    conn.commit()


# =====================================
# MARA ANALYTICS
# =====================================

with st.expander(
    "📊 MARA Analytics",
    expanded=False
):

    if st.session_state.notes:

        st.markdown(
            "### 📝 Generated Notes"
        )

        st.markdown(
            st.session_state.notes
        )

    if st.session_state.quiz:

        st.markdown(
            "### 🎯 Generated Quiz"
        )

        st.markdown(
            st.session_state.quiz
        )

    st.markdown(
        "### ⚡ Agent Activity"
    )

    st.info(
        f"Current Agent: "
        f"{st.session_state.current_agent}"
    )

    st.metric(
        "Agent Confidence",
        f"{st.session_state.confidence}%"
    )
    st.metric(
        "Average Processing Time",
        f"{st.session_state.processing_time} sec"
    )

    st.markdown(
    "### 📈 Agent Usage Statistics"
    )

    for agent_name, count in st.session_state.agent_usage.items():
        st.write(
            f"🤖 {agent_name}: {count}"
        )

    if st.session_state.agent_usage:

        st.bar_chart(
        st.session_state.agent_usage
    )
        
    if st.session_state.agent_usage:

        most_used = max(
        st.session_state.agent_usage,
        key=st.session_state.agent_usage.get
        )

        st.success(
            f"🏆 Most Used Agent: {most_used}"
        )

    st.markdown(
        "### ⚡ Agent Timeline"
    )

    if st.session_state.timeline:

        for step in st.session_state.timeline:

            st.write(step)

    else:

        st.write(
            "No activity yet."
        )

    st.markdown(
        "### ⭐ Feedback Analytics"
    )

    st.metric(
        "Total Feedback",
        len(
            st.session_state.feedback
        )
    )

    if st.session_state.feedback:

        helpful = (
            st.session_state.feedback.count(
                "👍 Helpful"
            )
        )

        st.metric(
            "Helpful Responses",
            helpful
        )
    
   
st.markdown("---")

st.caption(
    "© 2026 MARA AI Assistant | Built using LangGraph • Gemini • RAG • Streamlit"
)